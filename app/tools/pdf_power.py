from flask import Blueprint, render_template, request, send_file, jsonify, after_this_request, redirect, url_for
import os
import tempfile
import shutil
import traceback
import zipfile
import json
import textwrap
import uuid
from werkzeug.utils import secure_filename
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import Color
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from PyPDF2 import PdfReader, PdfWriter
import fitz  # PyMuPDF
import io
import pikepdf
import logging
from flask_login import login_required, current_user


pdf_power_bp = Blueprint("pdf_power", __name__, url_prefix="/tools/pdf_power")

# ログ設定
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ファイルサイズ制限（全機能共通）
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB

# 範囲コピー時の描画品質設定（高画質寄り）
COPY_REGION_BASE_DPI = 300
COPY_REGION_MAX_SCALE = 8.0

JAPANESE_FONT = "Helvetica"
REPORTLAB_JAPANESE_FONT = "Helvetica"

try:
    _reportlab_font_paths = [
        "C:/Windows/Fonts/msgothic.ttc",
        "C:/Windows/Fonts/meiryo.ttc",
        "C:/Windows/Fonts/YuGothM.ttc",
        "/System/Library/Fonts/NotoSansCJK.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf",
    ]
    for _font_path in _reportlab_font_paths:
        if os.path.exists(_font_path):
            try:
                pdfmetrics.registerFont(TTFont("NotoSansCJK", _font_path))
                JAPANESE_FONT = "NotoSansCJK"
                break
            except Exception:
                continue

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
        REPORTLAB_JAPANESE_FONT = "HeiseiKakuGo-W5"
    except Exception:
        REPORTLAB_JAPANESE_FONT = JAPANESE_FONT

    if REPORTLAB_JAPANESE_FONT == "Helvetica":
        logger.warning(
            "Japanese font was not found. PowerPDF text output will fall back to Helvetica."
        )
except Exception as e:
    logger.error(f"Japanese font registration error: {e}")
    logger.error(traceback.format_exc())


@pdf_power_bp.route("/", methods=["GET"])
@login_required
def pdf_power_home():
    return render_template(
        "pdf_power.html",
        initial_tab="control",
        password_mode="add",
        user_id=str(getattr(current_user, "username", "") or ""),
    )


@pdf_power_bp.route("/editor", methods=["GET"])
@login_required
def pdf_power_editor_window():
    return redirect(url_for("pdf_power.pdf_power_home", editor_window=1))


def _render_pdf_power_form_error(message, tab="control", status=400, **context):
    render_context = {
        "initial_tab": tab,
        "password_mode": context.get("password_mode", "add"),
        "password_error": context.get("password_error"),
        "user_id": str(getattr(current_user, "username", "") or ""),
    }
    render_context.update(context)
    return render_template("pdf_power.html", **render_context), status


def _parse_page_number(page_value, total_pages):
    """Convert 1-based page number to 0-based index."""
    try:
        page_number = int(page_value)
    except (TypeError, ValueError):
        raise ValueError("Page number must be an integer.")
    if page_number < 1 or page_number > total_pages:
        raise ValueError(f"Page number must be between 1 and {total_pages}.")
    return page_number - 1


def _parse_color_hex(color_value, default=(0, 0, 0)):
    """Convert '#RRGGBB' / 'RRGGBB' to PyMuPDF RGB tuple."""
    if not color_value:
        return default
    color = str(color_value).strip()
    if color.lower() in {"none", "transparent"}:
        return None
    if color.startswith("#"):
        color = color[1:]
    if len(color) != 6:
        return default
    try:
        r = int(color[0:2], 16) / 255
        g = int(color[2:4], 16) / 255
        b = int(color[4:6], 16) / 255
        return (r, g, b)
    except ValueError:
        return default


def _parse_page_ranges(pages_str, total_pages):
    """
    Parse page specification:
    - all
    - 1,3,5
    - 1-3,7,10-12
    Returns 0-based indices.
    """
    text = (pages_str or "").strip().lower()
    if text == "all":
        return list(range(total_pages))

    page_indices = []
    for raw_part in str(pages_str).split(","):
        part = raw_part.strip()
        if not part:
            continue
        if "-" in part:
            start_str, end_str = part.split("-", 1)
            start = int(start_str.strip())
            end = int(end_str.strip())
            if start > end:
                raise ValueError(f"Invalid range: {part}")
            page_indices.extend(range(start - 1, end))
        else:
            page_indices.append(int(part) - 1)

    if not page_indices:
        raise ValueError("Page range is empty.")

    for idx in page_indices:
        if idx < 0 or idx >= total_pages:
            raise ValueError(f"Out of range page number: {idx + 1}")

    return page_indices


_JAPANESE_FONT_CACHE = None
_FONT_REGISTERED = set()


def _resolve_japanese_font_file():
    """Resolve a usable Japanese font file path once and cache it."""
    global _JAPANESE_FONT_CACHE
    if _JAPANESE_FONT_CACHE is not None:
        return _JAPANESE_FONT_CACHE

    candidates = [
        "C:/Windows/Fonts/msgothic.ttc",
        "C:/Windows/Fonts/meiryo.ttc",
        "C:/Windows/Fonts/YuGothM.ttc",
        "/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf",
    ]

    for path in candidates:
        if os.path.exists(path):
            _JAPANESE_FONT_CACHE = path
            return path

    _JAPANESE_FONT_CACHE = ""
    return ""


def _is_non_ascii_text(text):
    return any(ord(ch) > 127 for ch in str(text))


def _insert_textbox_with_fallback(page, rect, text, font_size, color, font_name):
    """
    Insert text robustly:
    - ASCII: use requested font
    - Non-ASCII: try embedded Japanese font file, then CJK built-in font aliases, then requested font
    """
    if not _is_non_ascii_text(text):
        page.insert_textbox(
            rect,
            text,
            fontsize=font_size,
            fontname=font_name,
            color=color,
            overlay=True,
        )
        return

    font_file = _resolve_japanese_font_file()
    if font_file:
        try:
            alias = "jp_fallback"
            if alias not in _FONT_REGISTERED:
                page.insert_font(fontname=alias, fontfile=font_file)
                _FONT_REGISTERED.add(alias)
            page.insert_textbox(
                rect,
                text,
                fontsize=font_size,
                fontname=alias,
                color=color,
                overlay=True,
            )
            return
        except Exception:
            logger.warning("Japanese font embedding failed; trying CJK built-in font fallback.")

    for alias in ("japan", "japan-s", "japan-g"):
        try:
            page.insert_textbox(
                rect,
                text,
                fontsize=font_size,
                fontname=alias,
                color=color,
                overlay=True,
            )
            return
        except Exception:
            continue

    page.insert_textbox(
        rect,
        text,
        fontsize=font_size,
        fontname=font_name,
        color=color,
        overlay=True,
    )


def _dash_pattern(stroke_style):
    style = str(stroke_style or "solid").lower()
    if style == "dashed":
        return "[8 4] 0"
    if style == "dotted":
        return "[1 3] 0"
    return None


def _round_to_tenth(value):
    """Round numeric input to 1 decimal place for stable PDF rendering."""
    return round(float(value), 1)


def _to_pdf_point(page, x, y):
    """
    Convert viewer coordinates (PDF.js-rendered page space) to base PDF coordinates.
    The viewer follows page rotation metadata, while drawing APIs operate in base space.
    """
    xv = float(x)
    yv = float(y)
    rot = int(getattr(page, "rotation", 0)) % 360
    mbox = page.mediabox
    width = float(mbox.width)
    height = float(mbox.height)

    if rot == 90:
        return fitz.Point(yv, height - xv)
    if rot == 180:
        return fitz.Point(width - xv, height - yv)
    if rot == 270:
        return fitz.Point(width - yv, xv)
    return fitz.Point(xv, yv)


def _to_pdf_rect(page, x, y, w, h):
    # Map all corners to preserve correct bounds on rotated pages.
    x1 = float(x)
    y1 = float(y)
    x2 = x1 + float(w)
    y2 = y1 + float(h)
    corners = [
        _to_pdf_point(page, x1, y1),
        _to_pdf_point(page, x2, y1),
        _to_pdf_point(page, x1, y2),
        _to_pdf_point(page, x2, y2),
    ]
    xs = [p.x for p in corners]
    ys = [p.y for p in corners]
    rect = fitz.Rect(min(xs), min(ys), max(xs), max(ys))
    rect.normalize()
    return rect


def _normalize_edit_source_document(source_document):
    return "reference" if str(source_document or "").strip().lower() == "reference" else "edit"


def _collect_affected_pages_for_edit(operations, total_pages, reference_total_pages=0):
    """
    Collect 0-based page indices touched by edit operations.
    Used to normalize page rotation before drawing.
    """
    affected = {"edit": set(), "reference": set()}
    for idx, op in enumerate(operations, start=1):
        if not isinstance(op, dict):
            raise ValueError(f"Operation {idx} format is invalid.")

        op_type = op.get("type")
        if op_type in {"add_text", "add_shape", "add_image", "add_freehand"}:
            affected["edit"].add(_parse_page_number(op.get("page"), total_pages))
        elif op_type == "copy_region_paste":
            source_document = _normalize_edit_source_document(op.get("source_document"))
            if source_document == "reference":
                if reference_total_pages < 1:
                    raise ValueError("Reference PDF is required for reference copy operations.")
                affected["reference"].add(_parse_page_number(op.get("source_page"), reference_total_pages))
            else:
                affected["edit"].add(_parse_page_number(op.get("source_page"), total_pages))
            affected["edit"].add(_parse_page_number(op.get("target_page"), total_pages))

    return affected


def _is_pdf_filename(filename):
    return str(filename or "").lower().endswith(".pdf")


def _sanitize_upload_filename(upload_file, default_name="uploaded"):
    filename = secure_filename((getattr(upload_file, "filename", "") or "").strip())
    return filename or default_name


def _ensure_pdf_upload_filename(upload_file, default_name="uploaded.pdf"):
    filename = _sanitize_upload_filename(upload_file, default_name)
    if not _is_pdf_filename(filename):
        filename = f"{filename}.pdf"
    return filename


def _unique_temp_path(temp_dir, filename):
    safe_name = secure_filename(filename) or "file"
    stem, ext = os.path.splitext(safe_name)
    return os.path.join(temp_dir, f"{stem}_{uuid.uuid4().hex}{ext}")


def _send_path_bytes(file_path, download_name, mimetype=None):
    with open(file_path, "rb") as fh:
        file_data = io.BytesIO(fh.read())
    file_data.seek(0)
    return send_file(
        file_data,
        as_attachment=True,
        download_name=download_name,
        mimetype=mimetype,
    )


def _read_text_file(file_path):
    for encoding in ("utf-8-sig", "utf-8", "cp932", "shift_jis"):
        try:
            with open(file_path, "r", encoding=encoding) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _reportlab_font_name_for_text(text):
    return REPORTLAB_JAPANESE_FONT if _is_non_ascii_text(text) else "Helvetica"


def _wrap_reportlab_text(text, font_name, font_size, max_width, pdf_canvas):
    if text == "":
        return [""]

    wrapped_lines = []
    current = ""
    for ch in str(text):
        candidate = current + ch
        if not current or pdf_canvas.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            wrapped_lines.append(current)
            current = ch
    if current or not wrapped_lines:
        wrapped_lines.append(current)
    return wrapped_lines


def _create_text_pdf(content, output_path):
    page_width, page_height = letter
    margin = 72
    font_size = 12
    line_height = 18
    max_width = page_width - (margin * 2)

    pdf_canvas = canvas.Canvas(output_path, pagesize=letter)
    y = page_height - margin

    for raw_line in str(content).splitlines():
        font_name = _reportlab_font_name_for_text(raw_line)
        wrapped_lines = _wrap_reportlab_text(raw_line, font_name, font_size, max_width, pdf_canvas)
        for line in wrapped_lines:
            if y < margin:
                pdf_canvas.showPage()
                y = page_height - margin
            pdf_canvas.setFont(font_name, font_size)
            pdf_canvas.drawString(margin, y, line)
            y -= line_height
        if raw_line == "":
            y -= line_height / 2

    pdf_canvas.save()


def _create_image_pdf(image_path, output_path):
    with Image.open(image_path) as image:
        rgb_image = image.convert("RGB")
        rgb_image.save(output_path, "PDF")


def _create_watermark_overlay(page_rect, watermark_text):
    settings = {}
    if isinstance(watermark_text, dict):
        settings = watermark_text
        watermark_text = settings.get("text", "")

    text = str(watermark_text or "").strip()
    if not text:
        raise ValueError("Watermark text is required.")

    font_name = _reportlab_font_name_for_text(text)
    font_size = max(12, min(int(float(settings.get("font_size", 60) or 60)), 240))
    rotation = float(settings.get("rotation", 45) or 45)
    opacity = min(max(float(settings.get("opacity", 0.22) or 0.22), 0.05), 1.0)
    position = str(settings.get("position", "center") or "center").strip().lower()
    margin_x = max(0, float(settings.get("margin_x", 36) or 36))
    margin_y = max(0, float(settings.get("margin_y", 36) or 36))
    fill_rgb = _parse_color_hex(settings.get("color"), default=(0.8, 0.8, 0.8)) or (0.8, 0.8, 0.8)

    buffer = io.BytesIO()
    pdf_canvas = canvas.Canvas(buffer, pagesize=(page_rect.width, page_rect.height))
    pdf_canvas.setFillColor(Color(*fill_rgb))
    if hasattr(pdf_canvas, "setFillAlpha"):
        try:
            pdf_canvas.setFillAlpha(opacity)
        except Exception:
            pass
    pdf_canvas.setFont(font_name, font_size)

    text_width = pdf_canvas.stringWidth(text, font_name, font_size)
    half_width = text_width / 2
    half_height = font_size / 2
    x_positions = {
        "left": margin_x + half_width,
        "center": page_rect.width / 2,
        "right": max(page_rect.width / 2, page_rect.width - margin_x - half_width),
    }
    y_positions = {
        "top": max(page_rect.height / 2, page_rect.height - margin_y - half_height),
        "middle": page_rect.height / 2,
        "bottom": margin_y + half_height,
    }

    def draw_at(center_x, center_y):
        pdf_canvas.saveState()
        pdf_canvas.translate(center_x, center_y)
        pdf_canvas.rotate(rotation)
        pdf_canvas.drawCentredString(0, 0, text)
        pdf_canvas.restoreState()

    if position == "tile":
        step_x = max(text_width + 64, page_rect.width / 3)
        step_y = max(font_size * 2.8, page_rect.height / 3)
        start_x = margin_x + half_width
        start_y = margin_y + half_height
        max_x = max(start_x, page_rect.width - margin_x - half_width)
        max_y = max(start_y, page_rect.height - margin_y - half_height)
        y = start_y
        while y <= max_y:
            x = start_x
            while x <= max_x:
                draw_at(x, y)
                x += step_x
            y += step_y
    else:
        anchor_map = {
            "top-left": (x_positions["left"], y_positions["top"]),
            "top-center": (x_positions["center"], y_positions["top"]),
            "top-right": (x_positions["right"], y_positions["top"]),
            "middle-left": (x_positions["left"], y_positions["middle"]),
            "center": (x_positions["center"], y_positions["middle"]),
            "middle-right": (x_positions["right"], y_positions["middle"]),
            "bottom-left": (x_positions["left"], y_positions["bottom"]),
            "bottom-center": (x_positions["center"], y_positions["bottom"]),
            "bottom-right": (x_positions["right"], y_positions["bottom"]),
        }
        draw_at(*anchor_map.get(position, anchor_map["center"]))

    pdf_canvas.save()
    return buffer.getvalue()


def _normalize_pdf_download_name(filename, default_name="control_output.pdf"):
    name = str(filename or "").strip().replace("\\", "_").replace("/", "_")
    if not name:
        name = default_name
    if not name.lower().endswith(".pdf"):
        name = f"{name}.pdf"
    return name


def _coerce_bool(value, default=False):
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


def _coerce_int(value, default=0):
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _coerce_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _parse_control_plan(plan_raw):
    try:
        plan = json.loads(plan_raw or "{}")
    except json.JSONDecodeError as exc:
        raise ValueError("コントロール設定の読み込みに失敗しました。") from exc

    if not isinstance(plan, dict):
        raise ValueError("コントロール設定の形式が正しくありません。")
    if not isinstance(plan.get("documents"), list) or not plan["documents"]:
        raise ValueError("少なくとも1つのPDFを読み込んでください。")
    if not isinstance(plan.get("pages"), list) or not plan["pages"]:
        raise ValueError("出力対象のページが選択されていません。")

    return plan


def _open_control_source_pdf(file_path, password="", display_name="PDF"):
    doc = fitz.open(file_path)
    if not doc.needs_pass:
        return doc

    supplied_password = str(password or "")
    if not supplied_password:
        doc.close()
        raise ValueError(f"{display_name} はパスワードで保護されています。入力PDFのパスワードを設定してください。")

    if not doc.authenticate(supplied_password):
        doc.close()
        raise ValueError(f"{display_name} のパスワードが正しくありません。")

    return doc


def _build_control_metadata(selected_pages, docs_by_id, metadata_overrides):
    metadata = {}
    source_doc_ids = {page_plan.get("documentId") for page_plan in selected_pages}
    if len(source_doc_ids) == 1:
        source_doc = docs_by_id.get(next(iter(source_doc_ids)))
        if source_doc is not None:
            for key in ("title", "author", "subject", "keywords"):
                value = (source_doc.metadata or {}).get(key)
                if value:
                    metadata[key] = value

    overrides = metadata_overrides if isinstance(metadata_overrides, dict) else {}
    for key in ("title", "author", "subject", "keywords"):
        value = str(overrides.get(key, "") or "").strip()
        if value:
            metadata[key] = value

    return metadata


def _apply_control_watermark(doc, watermark_settings):
    if not isinstance(watermark_settings, dict) or not _coerce_bool(watermark_settings.get("enabled")):
        return

    text = str(watermark_settings.get("text", "") or "").strip()
    if not text or len(doc) == 0:
        return

    page_spec = str(watermark_settings.get("pages", "all") or "all").strip()
    target_pages = set(_parse_page_ranges(page_spec, len(doc)))

    overlay_settings = {
        "text": text,
        "font_size": _coerce_int(watermark_settings.get("font_size"), 60),
        "rotation": _coerce_float(watermark_settings.get("rotation"), 45.0),
        "opacity": _coerce_float(watermark_settings.get("opacity"), 0.22),
        "position": str(watermark_settings.get("position", "center") or "center"),
        "margin_x": _coerce_float(watermark_settings.get("margin_x"), 36.0),
        "margin_y": _coerce_float(watermark_settings.get("margin_y"), 36.0),
        "color": watermark_settings.get("color", "#c0c7d1"),
    }

    for page_index in target_pages:
        page = doc[page_index]
        overlay_pdf = fitz.open(
            stream=_create_watermark_overlay(page.rect, overlay_settings),
            filetype="pdf",
        )
        try:
            page.show_pdf_page(page.rect, overlay_pdf, 0, overlay=True)
        finally:
            overlay_pdf.close()


# ========================================
# 1. PDF変換機能
# ========================================

def convert_from_pdf():
    """
    PDF→画像/TXT/DOCXの逆変換
    """
    temp_dir = None
    try:
        file = request.files.get("pdf_file")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "PDFファイルのみ対応しています"}), 400

        output_format = request.form.get("output_format", "png")
        if output_format not in ["png", "jpg", "txt", "docx"]:
            return jsonify({"error": "対応していない出力形式です"}), 400

        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, filename)
        file.save(input_path)

        # ファイルサイズチェック
        if os.path.getsize(input_path) > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        # PDFの有効性確認
        try:
            doc = fitz.open(input_path)
            total_pages = len(doc)
            if total_pages == 0:
                return jsonify({"error": "PDFファイルが空です"}), 400
        except Exception as e:
            return jsonify({"error": f"PDFファイルが破損しています: {str(e)}"}), 400

        # 出力形式に応じた変換処理
        if output_format in ["png", "jpg"]:
            # PDF→画像変換
            output_files = []
            for page_num in range(total_pages):
                page = doc[page_num]
                # 高解像度でレンダリング（DPI 150）
                zoom = 2.0  # 150 DPI相当
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)

                image_ext = output_format
                image_filename = f"page_{page_num + 1:03d}.{image_ext}"
                image_path = os.path.join(temp_dir, image_filename)

                if output_format == "png":
                    pix.save(image_path)
                else:  # jpg
                    # PNGで保存してからJPEGに変換
                    temp_png = image_path + ".png"
                    pix.save(temp_png)
                    img = Image.open(temp_png)
                    img = img.convert("RGB")
                    img.save(image_path, "JPEG", quality=95)
                    os.remove(temp_png)

                output_files.append({"path": image_path, "name": image_filename})

            doc.close()

            # 複数ファイルをZIPにまとめる
            if len(output_files) > 1:
                zip_path = os.path.join(temp_dir, f"pdf_to_{output_format}.zip")
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file_info in output_files:
                        zipf.write(file_info["path"], file_info["name"])

                @after_this_request
                def cleanup(response):
                    try:
                        if temp_dir and os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
                    except Exception as e:
                        logger.error(f"一時ディレクトリの削除に失敗: {e}")
                    return response

                return _send_path_bytes(zip_path, f"pdf_to_{output_format}.zip", "application/zip")
            else:
                # 1ページのみの場合は単一ファイルを返す
                @after_this_request
                def cleanup(response):
                    try:
                        if temp_dir and os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
                    except Exception as e:
                        logger.error(f"一時ディレクトリの削除に失敗: {e}")
                    return response

                return _send_path_bytes(output_files[0]["path"], output_files[0]["name"])

        elif output_format == "txt":
            # PDF→TXT変換
            all_text = []
            for page_num in range(total_pages):
                page = doc[page_num]
                page_text = page.get_text()
                all_text.append(f"========== ページ {page_num + 1} ==========\n")
                all_text.append(page_text)
                all_text.append("\n\n")

            doc.close()

            txt_path = os.path.join(temp_dir, "converted_output.txt")
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write("".join(all_text))

            @after_this_request
            def cleanup(response):
                try:
                    if temp_dir and os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.error(f"一時ディレクトリの削除に失敗: {e}")
                return response

            return _send_path_bytes(txt_path, "converted_output.txt", "text/plain; charset=utf-8")

        elif output_format == "docx":
            # PDF→DOCX変換
            docx_doc = Document()

            for page_num in range(total_pages):
                page = doc[page_num]
                page_text = page.get_text()

                # ページ見出しを追加
                docx_doc.add_heading(f"ページ {page_num + 1}", level=2)

                # テキストを段落として追加
                paragraphs = page_text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        docx_doc.add_paragraph(para_text)

                # ページの区切り（最後のページ以外）
                if page_num < total_pages - 1:
                    docx_doc.add_page_break()

            doc.close()

            docx_path = os.path.join(temp_dir, "converted_output.docx")
            docx_doc.save(docx_path)

            @after_this_request
            def cleanup(response):
                try:
                    if temp_dir and os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.error(f"一時ディレクトリの削除に失敗: {e}")
                return response

            return _send_path_bytes(
                docx_path,
                "converted_output.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    except Exception as e:
        logger.error(f"PDF変換処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"変換処理中にエラーが発生しました: {str(e)}"}), 500
    finally:
        pass


@pdf_power_bp.route("/convert", methods=["POST"])
@login_required
def convert_pdf():
    temp_dir = None
    output_path = None

    try:
        # 変換方向の取得
        direction = request.form.get("direction", "to_pdf")

        # PDF→ファイル変換の場合
        if direction == "from_pdf":
            return convert_from_pdf()

        files = request.files.getlist("files")
        if not files or not files[0].filename:
            return jsonify({"error": "ファイルがアップロードされていません"}), 400

        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, "converted_output.pdf")
        part_paths = []

        for index, file in enumerate(files, start=1):
            if not file.filename:
                continue

            filename = _sanitize_upload_filename(file, f"upload_{index}")
            ext = os.path.splitext(filename)[1].lower()
            file_path = _unique_temp_path(temp_dir, filename)
            file.save(file_path)

            if os.path.getsize(file_path) > MAX_FILE_SIZE:
                return jsonify({"error": f"ファイルサイズが大きすぎます（100MB以下にしてください）: {filename}"}), 400

            part_path = os.path.join(temp_dir, f"convert_part_{index:03d}.pdf")

            if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif"]:
                try:
                    _create_image_pdf(file_path, part_path)
                except Exception as e:
                    logger.error(f"画像の変換に失敗しました: {e}")
                    logger.error(traceback.format_exc())
                    return jsonify({"error": f"画像の変換に失敗しました: {filename}"}), 400
            elif ext == ".txt":
                try:
                    _create_text_pdf(_read_text_file(file_path), part_path)
                except Exception as e:
                    logger.error(f"テキストファイルの変換に失敗しました: {e}")
                    logger.error(traceback.format_exc())
                    return jsonify({"error": f"テキストファイルの変換に失敗しました: {str(e)}"}), 400
            elif ext == ".docx":
                try:
                    word_doc = Document(file_path)
                    content = "\n".join(para.text for para in word_doc.paragraphs)
                    _create_text_pdf(content, part_path)
                except Exception as e:
                    logger.error(f"Wordファイルの変換に失敗しました: {e}")
                    logger.error(traceback.format_exc())
                    return jsonify({"error": f"Wordファイルの変換に失敗しました: {str(e)}"}), 400
            else:
                ext_label = ext[1:] if ext.startswith(".") else ext
                return jsonify({"error": f"対応していないファイル形式です: {ext_label}. PNG, JPG, JPEG, BMP, GIF, TXT, DOCXのみ対応しています。"}), 400

            part_paths.append(part_path)

        if not part_paths:
            return jsonify({"error": "変換するファイルが見つかりませんでした"}), 400

        writer = PdfWriter()
        for part_path in part_paths:
            reader = PdfReader(part_path)
            for page in reader.pages:
                writer.add_page(page)

        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)

        @after_this_request
        def cleanup(response):
            try:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except Exception as e:
                logger.error(f"一時ディレクトリの削除に失敗: {e}")
            return response

        return _send_path_bytes(output_path, "converted_output.pdf", "application/pdf")

    except Exception as e:
        logger.error(f"変換処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"変換処理中にエラーが発生しました: {str(e)}"}), 500
    finally:
        # after_this_requestがない場合のフォールバック
        pass


# ========================================
# 2. PDF分割・結合機能
# ========================================
@pdf_power_bp.route("/split_merge", methods=["POST"])
@login_required
def split_or_merge_pdf():
    temp_dir = None
    try:
        mode = request.form.get("mode")
        if mode not in ["split", "merge"]:
            return jsonify({"error": "無効な操作です"}), 400

        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, "output.pdf")

        files = request.files.getlist("pdfs")
        if not files or not files[0].filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        file_paths = []
        for file in files:
            if not file.filename:
                continue
            filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
            if not _is_probably_pdf_upload(file):
                return jsonify({"error": f"PDFファイルのみ対応しています: {filename}"}), 400

            file_path = _unique_temp_path(temp_dir, filename)
            file.save(file_path)

            # ファイルサイズチェック
            if os.path.getsize(file_path) > MAX_FILE_SIZE:
                return jsonify({"error": f"ファイルサイズが大きすぎます（100MB以下にしてください）: {filename}"}), 400

            file_paths.append(file_path)

        if not file_paths:
            return jsonify({"error": "有効なPDFファイルが見つかりませんでした"}), 400

        # 分割
        if mode == "split":
            if len(file_paths) != 1:
                return jsonify({"error": "分割は1つのPDFファイルのみ対応しています"}), 400

            page_range = request.form.get("range")
            if not page_range:
                return jsonify({"error": "ページ範囲を指定してください（例: 1-5）"}), 400

            try:
                reader = PdfReader(file_paths[0])
                total_pages = len(reader.pages)
                pages_to_extract = _parse_page_ranges(page_range, total_pages)
            except ValueError:
                return jsonify({"error": "ページ範囲の指定が正しくありません（例: 1-5）"}), 400

            try:
                writer = PdfWriter()
                for page_num in pages_to_extract:
                    writer.add_page(reader.pages[page_num])

                with open(output_path, "wb") as output_pdf:
                    writer.write(output_pdf)

                @after_this_request
                def cleanup(response):
                    try:
                        if temp_dir and os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
                    except Exception as e:
                        logger.error(f"一時ディレクトリの削除に失敗: {e}")
                    return response

                return _send_path_bytes(output_path, "split_output.pdf", "application/pdf")

            except Exception as e:
                logger.error(f"PDF分割に失敗しました: {e}")
                logger.error(traceback.format_exc())
                return jsonify({"error": f"PDF分割に失敗しました: {str(e)}"}), 500

        # 結合
        elif mode == "merge":
            if len(file_paths) < 2:
                return jsonify({"error": "結合には2つ以上のPDFファイルが必要です"}), 400

            try:
                writer = PdfWriter()
                for file_path in file_paths:
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        writer.add_page(page)

                with open(output_path, "wb") as output_pdf:
                    writer.write(output_pdf)

                @after_this_request
                def cleanup(response):
                    try:
                        if temp_dir and os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
                    except Exception as e:
                        logger.error(f"一時ディレクトリの削除に失敗: {e}")
                    return response

                return _send_path_bytes(output_path, "merged_output.pdf", "application/pdf")

            except Exception as e:
                logger.error(f"PDF結合に失敗しました: {e}")
                logger.error(traceback.format_exc())
                return jsonify({"error": f"PDF結合に失敗しました: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"PDF操作中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"PDF操作中にエラーが発生しました: {str(e)}"}), 500
    finally:
        pass


# ========================================
# 3. PDF圧縮機能（改善版）
# ========================================
@pdf_power_bp.route("/control_process", methods=["POST"])
@login_required
def process_control_pdf():
    temp_dir = None
    output_doc = None
    source_docs = []

    try:
        plan = _parse_control_plan(request.form.get("plan"))
        uploads = request.files.getlist("pdfs")
        if not uploads:
            return jsonify({"error": "PDFファイルがアップロードされていません。"}), 400

        documents_plan = plan.get("documents", [])
        if len(uploads) < len(documents_plan):
            return jsonify({"error": "アップロードされたPDF数が設定内容と一致しません。"}), 400

        temp_dir = tempfile.mkdtemp()
        docs_by_id = {}

        for fallback_index, doc_plan in enumerate(documents_plan):
            if not isinstance(doc_plan, dict):
                return jsonify({"error": "PDF設定の形式が正しくありません。"}), 400

            document_id = str(doc_plan.get("id", "") or "").strip()
            if not document_id:
                return jsonify({"error": "PDF設定にIDがありません。"}), 400

            upload_index = _coerce_int(doc_plan.get("uploadIndex"), fallback_index)
            if upload_index < 0 or upload_index >= len(uploads):
                return jsonify({"error": "アップロードPDFの参照位置が不正です。"}), 400

            upload = uploads[upload_index]
            if not upload or not upload.filename:
                return jsonify({"error": "PDFファイルの読み込みに失敗しました。"}), 400
            if not _is_probably_pdf_upload(upload):
                return jsonify({"error": f"PDFファイルのみ対応しています: {upload.filename}"}), 400

            filename = _ensure_pdf_upload_filename(upload, f"source_{fallback_index + 1}.pdf")
            input_path = _unique_temp_path(temp_dir, filename)
            upload.save(input_path)

            if os.path.getsize(input_path) > MAX_FILE_SIZE:
                return jsonify({"error": f"ファイルサイズが大きすぎます: {filename}"}), 400

            display_name = str(doc_plan.get("name") or upload.filename or f"PDF {fallback_index + 1}")
            password = str(doc_plan.get("password", "") or "")
            source_doc = _open_control_source_pdf(input_path, password=password, display_name=display_name)
            source_docs.append(source_doc)
            docs_by_id[document_id] = source_doc

        selected_pages = plan.get("pages", [])
        output_doc = fitz.open()
        for page_order, page_plan in enumerate(selected_pages, start=1):
            if not isinstance(page_plan, dict):
                return jsonify({"error": f"{page_order}番目のページ設定が不正です。"}), 400

            document_id = str(page_plan.get("documentId", "") or "").strip()
            source_doc = docs_by_id.get(document_id)
            if source_doc is None:
                return jsonify({"error": f"{page_order}番目のページで参照しているPDFが見つかりません。"}), 400

            source_page_index = _coerce_int(page_plan.get("sourcePageIndex"), -1)
            if source_page_index < 0 or source_page_index >= len(source_doc):
                return jsonify({"error": f"{page_order}番目のページ番号が不正です。"}), 400

            output_doc.insert_pdf(source_doc, from_page=source_page_index, to_page=source_page_index)
            extra_rotation = _coerce_int(page_plan.get("rotation"), 0) % 360
            if extra_rotation:
                inserted_page = output_doc[-1]
                current_rotation = int(getattr(inserted_page, "rotation", 0)) % 360
                inserted_page.set_rotation((current_rotation + extra_rotation) % 360)

        if len(output_doc) == 0:
            return jsonify({"error": "出力対象のページがありません。"}), 400

        metadata = _build_control_metadata(selected_pages, docs_by_id, plan.get("metadata"))
        if metadata:
            output_doc.set_metadata(metadata)

        try:
            _apply_control_watermark(output_doc, plan.get("watermark"))
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400

        output_password_settings = plan.get("outputPassword") if isinstance(plan.get("outputPassword"), dict) else {}
        output_password_enabled = _coerce_bool(output_password_settings.get("enabled"))
        output_password = str(output_password_settings.get("password", "") or "")
        if output_password_enabled and not output_password:
            return jsonify({"error": "出力パスワードを設定する場合はパスワードを入力してください。"}), 400

        output_path = os.path.join(temp_dir, "control_output.pdf")
        save_kwargs = {"garbage": 3, "deflate": True}
        if output_password_enabled:
            save_kwargs.update(
                {
                    "encryption": fitz.PDF_ENCRYPT_AES_256,
                    "owner_pw": output_password,
                    "user_pw": output_password,
                }
            )
        output_doc.save(output_path, **save_kwargs)
        output_doc.close()
        output_doc = None

        download_name = _normalize_pdf_download_name(plan.get("outputName"), "control_output.pdf")

        @after_this_request
        def cleanup(response):
            try:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except Exception as e:
                logger.error(f"一時ディレクトリの削除に失敗: {e}")
            return response

        return _send_path_bytes(output_path, download_name, "application/pdf")

    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as e:
        logger.error(f"コントロール処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"コントロール処理中にエラーが発生しました: {str(e)}"}), 500
    finally:
        if output_doc is not None:
            output_doc.close()
        for source_doc in source_docs:
            try:
                source_doc.close()
            except Exception:
                pass


@pdf_power_bp.route("/compress", methods=["POST"])
@login_required
def compress_pdf():
    temp_dir = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = secure_filename(file.filename or "")
        if not filename:
            filename = "uploaded.pdf"

        compression_level = request.form.get("compression_level", "medium")

        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, filename)
        output_path = os.path.join(temp_dir, "compressed_output.pdf")
        file.save(input_path)

        # ファイルサイズチェック
        original_size = os.path.getsize(input_path)
        if original_size > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        # PDFの有効性チェック
        try:
            test_doc = fitz.open(input_path)
            if test_doc.page_count == 0:
                return jsonify({"error": "PDFファイルが空です"}), 400
            test_doc.close()
        except Exception as e:
            return jsonify({"error": f"PDFファイルが破損しています: {str(e)}"}), 400

        try:
            # 改善された圧縮処理（pikepdfベース）
            compressed_size = compress_pdf_advanced(
                input_path, output_path, compression_level
            )

            if compressed_size == 0:
                return jsonify({"error": "PDF圧縮に失敗しました"}), 500

            # 圧縮後にサイズが増加した場合は元のファイルを返す
            if compressed_size >= original_size:
                logger.info(f"圧縮後のサイズが元のサイズ以上のため、元のファイルを返します: {original_size} -> {compressed_size} bytes")
                send_path = input_path
                compression_ratio = 0.0
            else:
                send_path = output_path
                compression_ratio = (1 - compressed_size / original_size) * 100

            logger.info(f"PDF圧縮完了: {original_size} -> {os.path.getsize(send_path)} bytes ({compression_ratio:.1f}% 削減)")

            # ファイルをメモリに読み込んでから返す（Windowsでのファイルロック回避）
            with open(send_path, 'rb') as f:
                file_data = io.BytesIO(f.read())

            response = send_file(file_data, as_attachment=True, download_name="compressed_output.pdf", mimetype="application/pdf")
            response.headers['X-Original-Size'] = str(original_size)
            response.headers['X-Compressed-Size'] = str(os.path.getsize(send_path))
            response.headers['X-Compression-Ratio'] = f"{compression_ratio:.1f}"
            return response

        except Exception as e:
            logger.error(f"PDF圧縮に失敗しました: {e}")
            logger.error(traceback.format_exc())
            return jsonify({"error": f"PDF圧縮に失敗しました: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"圧縮処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"圧縮処理中にエラーが発生しました: {str(e)}"}), 500
    finally:
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                logger.error(f"一時ディレクトリの削除に失敗: {e}")


def _is_scan_page(page, threshold=0.8):
    """ページがスキャン画像かどうかを判定する。
    ページ面積の80%以上を画像が占めている場合にTrueを返す。
    """
    try:
        page_area = abs(page.rect)
        if page_area == 0:
            return False

        image_list = page.get_images(full=True)
        if not image_list:
            return False

        total_image_area = 0
        for img in image_list:
            xref = img[0]
            try:
                rects = page.get_image_rects(xref)
                for rect in rects:
                    if rect.is_empty or rect.is_infinite:
                        continue
                    total_image_area += abs(rect)
            except Exception:
                continue

        return (total_image_area / page_area) >= threshold
    except Exception:
        return False


def compress_pdf_advanced(input_path, output_path, compression_level="medium"):
    """
    PDF圧縮処理（pikepdf + PyMuPDFのハイブリッド）
    Step 1: PyMuPDFで構造最適化（メタデータ削除、ガベージコレクション）
            スキャンページは指定DPIで再レンダリング
    Step 2: pikepdfで非スキャンページの画像圧縮と追加最適化
    """
    try:
        # 圧縮レベル設定
        compression_settings = {
            "low": {"quality": 85, "max_size": (1200, 1200), "dpi": 150},
            "medium": {"quality": 70, "max_size": (1000, 1000), "dpi": 120},
            "high": {"quality": 50, "max_size": (800, 800), "dpi": 96},
            "maximum": {"quality": 30, "max_size": (600, 600), "dpi": 72}
        }

        settings = compression_settings.get(compression_level, compression_settings["medium"])

        # Step 1: PyMuPDFで構造最適化 + スキャンページ再レンダリング
        doc = fitz.open(input_path)

        # スキャンページの判定
        scan_page_indices = set()
        for page_num in range(len(doc)):
            if _is_scan_page(doc[page_num]):
                scan_page_indices.add(page_num)

        if scan_page_indices:
            logger.info(
                f"スキャンページ検出: {sorted(p + 1 for p in scan_page_indices)} "
                f"({len(scan_page_indices)}/{len(doc)}ページ)"
            )
            # スキャンページを再レンダリングした新しいドキュメントを作成
            new_doc = fitz.open()
            for page_num in range(len(doc)):
                if page_num in scan_page_indices:
                    page = doc[page_num]
                    # 指定DPIでページをラスタライズ
                    pix = page.get_pixmap(dpi=settings["dpi"], alpha=False)
                    img_data = pix.tobytes("jpeg", jpg_quality=settings["quality"])
                    # 元のページと同じサイズで新しいページを作成
                    new_page = new_doc.new_page(
                        width=page.rect.width, height=page.rect.height
                    )
                    new_page.insert_image(new_page.rect, stream=img_data)
                else:
                    new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)

            new_doc.set_metadata({})
            temp_output = output_path + ".tmp"
            new_doc.save(temp_output, garbage=4, deflate=True, clean=True)
            new_doc.close()
            doc.close()
        else:
            # スキャンページなし - 構造最適化のみ
            doc.set_metadata({})
            temp_output = output_path + ".tmp"
            doc.save(temp_output, garbage=4, deflate=True, clean=True)
            doc.close()

        # Step 2: pikepdfで非スキャンページの画像圧縮と追加最適化
        try:
            with pikepdf.open(temp_output) as pdf:
                # 各ページから未使用リソースを削除
                for page in pdf.pages:
                    try:
                        page.remove_unreferenced_resources()
                    except Exception as e:
                        logger.warning(f"未使用リソース削除スキップ: {e}")

                # 非スキャンページの画像を再圧縮
                for page_num, page in enumerate(pdf.pages):
                    if page_num in scan_page_indices:
                        continue  # 再レンダリング済みページはスキップ

                    try:
                        image_keys = list(page.images.keys())
                    except Exception:
                        continue

                    for image_key in image_keys:
                        try:
                            raw_image = page.images[image_key]
                            pdf_image = pikepdf.PdfImage(raw_image)
                            pil_image = pdf_image.as_pil_image()

                            # 画像圧縮
                            if pil_image.size[0] > settings["max_size"][0] or pil_image.size[1] > settings["max_size"][1]:
                                pil_image.thumbnail(settings["max_size"], Image.Resampling.LANCZOS)

                            # JPEG形式で保存
                            img_buffer = io.BytesIO()
                            if pil_image.mode in ("RGBA", "LA", "P"):
                                background = Image.new("RGB", pil_image.size, (255, 255, 255))
                                if pil_image.mode == "P":
                                    pil_image = pil_image.convert("RGBA")
                                if pil_image.mode == "RGBA":
                                    background.paste(pil_image, mask=pil_image.split()[-1])
                                pil_image = background
                            elif pil_image.mode != "RGB":
                                pil_image = pil_image.convert("RGB")

                            pil_image.save(img_buffer, format="JPEG", quality=settings["quality"], optimize=True)

                            # 圧縮画像を置換
                            img_buffer.seek(0)
                            raw_image.write(img_buffer.read(), filter=pikepdf.Name.DCTDecode)
                        except Exception as e:
                            logger.warning(f"pikepdf画像圧縮スキップ ({image_key}): {e}")
                            continue

                # 最終保存
                pdf.save(
                    output_path,
                    compress_streams=True,
                    stream_decode_level=pikepdf.StreamDecodeLevel.generalized,
                    object_stream_mode=pikepdf.ObjectStreamMode.generate,
                    linearize=True
                )
        except Exception as e:
            logger.warning(f"pikepdf最適化でエラー: {e}")
            logger.warning(traceback.format_exc())
            # pikepdfが失敗した場合は一時ファイルをコピー
            shutil.copy(temp_output, output_path)

        # 一時ファイル削除
        if os.path.exists(temp_output):
            os.remove(temp_output)

        return os.path.getsize(output_path)

    except Exception as e:
        logger.error(f"PDF圧縮処理に失敗しました: {e}")
        logger.error(traceback.format_exc())
        return 0


# ========================================
# 4. テキスト抽出機能（画像抽出対応）
# ========================================
@pdf_power_bp.route("/extract", methods=["POST"])
@login_required
def extract_text():
    """
    PDFテキスト抽出機能（画像抽出にも対応）
    """
    temp_dir = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "PDFファイルのみ対応しています"}), 400

        keyword = request.form.get("keyword", "").strip()
        extract_images = request.form.get("extract_images", "false").lower() == "true"

        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, filename)
        file.save(input_path)

        # ファイルサイズチェック
        file_size = os.path.getsize(input_path)
        if file_size > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        # PDFの有効性確認
        if not _validate_pdf_file(input_path):
            return jsonify({"error": "PDFファイルが破損しているか、読み取れません"}), 400

        # テキスト抽出
        extraction_result = _extract_text_from_pdf(input_path, keyword, extract_images, temp_dir)

        if extraction_result["success"]:
            output_files = _create_output_files(extraction_result, temp_dir)

            # 単一ファイルまたはZIPファイルの返却
            if len(output_files) == 1:
                @after_this_request
                def cleanup(response):
                    try:
                        if temp_dir and os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
                    except Exception as e:
                        logger.error(f"一時ディレクトリの削除に失敗: {e}")
                    return response

                return _send_path_bytes(output_files[0]["path"], output_files[0]["name"])
            else:
                zip_path = _create_zip_file(output_files, temp_dir)

                @after_this_request
                def cleanup(response):
                    try:
                        if temp_dir and os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
                    except Exception as e:
                        logger.error(f"一時ディレクトリの削除に失敗: {e}")
                    return response

                return _send_path_bytes(zip_path, "extracted_content.zip", "application/zip")
        else:
            return jsonify({"error": extraction_result["error"]}), 500

    except Exception as e:
        logger.error(f"テキスト抽出処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"予期しないエラーが発生しました: {str(e)}"}), 500
    finally:
        pass


def _temp_input_pdf_path(temp_dir):
    """PyMuPDFの保存先と衝突しない一時入力PDFパスを返す。"""
    return os.path.join(temp_dir, "uploaded_input.pdf")

def _validate_pdf_file(file_path):
    """PDFファイルの有効性を確認"""
    try:
        doc = fitz.open(file_path)
        if doc.page_count == 0:
            doc.close()
            return False
        doc.close()

        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            if len(reader.pages) == 0:
                return False

        return True
    except Exception as e:
        logger.error(f"PDFファイル検証エラー: {e}")
        return False


def _is_probably_pdf_upload(upload_file):
    """拡張子が無くても、ヘッダとMIMEからPDFアップロードを判定する。"""
    if not upload_file:
        return False

    filename = secure_filename(upload_file.filename or "")
    mimetype = (upload_file.mimetype or "").lower()

    stream = upload_file.stream
    current_pos = stream.tell()
    try:
        stream.seek(0)
        header = stream.read(5)
    finally:
        stream.seek(current_pos)

    if header == b"%PDF-":
        return True
    if filename.lower().endswith(".pdf"):
        return True
    if mimetype in {"application/pdf", "application/x-pdf"}:
        return True

    return False


def _extract_text_from_pdf(file_path, keyword, extract_images, temp_dir):
    """PDFからテキストを抽出"""
    try:
        result = {
            "success": False,
            "text": "",
            "keyword_matches": [],
            "images": [],
            "statistics": {},
            "error": None
        }

        pymupdf_result = _extract_with_pymupdf(file_path, keyword, extract_images, temp_dir)
        pypdf2_result = _extract_with_pypdf2(file_path, keyword)

        if pymupdf_result["success"] and len(pymupdf_result["text"]) > 0:
            result = pymupdf_result
        elif pypdf2_result["success"] and len(pypdf2_result["text"]) > 0:
            result = pypdf2_result
            if extract_images and pymupdf_result["images"]:
                result["images"] = pymupdf_result["images"]
        else:
            result["error"] = "どの手法でもテキストを抽出できませんでした"
            return result

        result["statistics"] = _calculate_statistics(result, file_path)
        result["success"] = True

        return result

    except Exception as e:
        logger.error(f"テキスト抽出処理でエラー: {e}")
        logger.error(traceback.format_exc())
        return {
            "success": False,
            "error": f"テキスト抽出処理でエラー: {str(e)}"
        }


def _extract_with_pymupdf(file_path, keyword, extract_images, temp_dir):
    """PyMuPDFを使用したテキスト抽出"""
    result = {
        "success": False,
        "text": "",
        "keyword_matches": [],
        "images": [],
        "method": "PyMuPDF"
    }

    try:
        doc = fitz.open(file_path)
        all_text = []
        keyword_matches = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_number = page_num + 1

            page_text = ""
            try:
                page_text = page.get_text()
            except Exception:
                try:
                    page_text = page.get_text("text")
                except Exception:
                    pass

            if len(page_text.strip()) < 10:
                try:
                    text_dict = page.get_text("dict")
                    extracted_text = ""
                    for block in text_dict.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line.get("spans", []):
                                    extracted_text += span.get("text", "")
                                extracted_text += "\n"
                    page_text = extracted_text
                except Exception:
                    pass

            if not page_text.strip():
                page_text = f"[ページ {page_number}: テキストが検出されませんでした]\n"

            if keyword:
                keyword_matches.extend(_find_keyword_matches(page_text, keyword, page_number))

            if extract_images:
                try:
                    image_list = page.get_images(full=True)
                    for img_index, img in enumerate(image_list):
                        try:
                            image_path = _extract_image(doc, img, page_number, img_index, temp_dir)
                            if image_path:
                                result["images"].append({
                                    "page": page_number,
                                    "index": img_index + 1,
                                    "path": image_path,
                                    "filename": os.path.basename(image_path)
                                })
                        except Exception as e:
                            logger.warning(f"画像抽出エラー (ページ {page_number}, 画像 {img_index + 1}): {e}")
                except Exception as e:
                    logger.warning(f"ページ {page_number} の画像処理エラー: {e}")

            all_text.append(f"=== ページ {page_number} ===\n{page_text.strip()}\n")

        doc.close()

        result["text"] = "\n".join(all_text)
        result["keyword_matches"] = keyword_matches
        result["success"] = True

        return result

    except Exception as e:
        logger.error(f"PyMuPDF抽出エラー: {e}")
        result["error"] = f"PyMuPDF抽出エラー: {str(e)}"
        return result


def _extract_with_pypdf2(file_path, keyword):
    """PyPDF2を使用したテキスト抽出"""
    result = {
        "success": False,
        "text": "",
        "keyword_matches": [],
        "images": [],
        "method": "PyPDF2"
    }

    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            all_text = []
            keyword_matches = []

            for page_num, page in enumerate(reader.pages):
                page_number = page_num + 1

                try:
                    page_text = page.extract_text()

                    if not page_text.strip():
                        page_text = f"[ページ {page_number}: テキストが検出されませんでした]\n"

                    if keyword:
                        keyword_matches.extend(_find_keyword_matches(page_text, keyword, page_number))

                    all_text.append(f"=== ページ {page_number} ===\n{page_text.strip()}\n")

                except Exception as e:
                    logger.warning(f"ページ {page_number} の処理エラー: {e}")
                    all_text.append(f"=== ページ {page_number} ===\n[処理エラー: {str(e)}]\n")

            result["text"] = "\n".join(all_text)
            result["keyword_matches"] = keyword_matches
            result["success"] = True

        return result

    except Exception as e:
        logger.error(f"PyPDF2抽出エラー: {e}")
        result["error"] = f"PyPDF2抽出エラー: {str(e)}"
        return result


def _find_keyword_matches(text, keyword, page_number):
    """テキスト内のキーワード検索"""
    matches = []
    if not keyword or not text:
        return matches

    lines = text.split('\n')
    keyword_lower = keyword.lower()

    for line_num, line in enumerate(lines):
        if keyword_lower in line.lower():
            context_start = max(0, line_num - 2)
            context_end = min(len(lines), line_num + 3)
            context = '\n'.join(lines[context_start:context_end])

            matches.append({
                'page': page_number,
                'line_number': line_num + 1,
                'line': line.strip(),
                'context': context.strip()
            })

    return matches


def _extract_image(doc, img, page_number, img_index, temp_dir):
    """PDFから画像を抽出"""
    try:
        xref = img[0]
        base_image = doc.extract_image(xref)
        image_bytes = base_image["image"]
        image_ext = base_image["ext"]

        image_filename = f"page_{page_number:03d}_image_{img_index + 1:03d}.{image_ext}"
        image_path = os.path.join(temp_dir, image_filename)

        with open(image_path, "wb") as img_file:
            img_file.write(image_bytes)

        return image_path

    except Exception as e:
        logger.warning(f"画像抽出エラー: {e}")
        return None


def _calculate_statistics(result, file_path):
    """抽出結果の統計情報を計算"""
    try:
        doc = fitz.open(file_path)
        stats = {
            "total_pages": len(doc),
            "total_characters": len(result["text"]),
            "total_words": len(result["text"].split()),
            "total_lines": len(result["text"].split('\n')),
            "file_size": os.path.getsize(file_path),
            "extraction_method": result.get("method", "Unknown"),
            "keyword_matches": len(result["keyword_matches"]),
            "extracted_images": len(result["images"])
        }
        doc.close()
        return stats
    except Exception as e:
        logger.error(f"統計計算エラー: {e}")
        return {}


def _create_output_files(result, temp_dir):
    """出力ファイルの作成"""
    output_files = []

    text_content = _format_extracted_text(result)
    text_path = os.path.join(temp_dir, "extracted_text.txt")

    with open(text_path, 'w', encoding='utf-8') as f:
        f.write(text_content)

    output_files.append({
        "path": text_path,
        "name": "extracted_text.txt",
        "type": "text"
    })

    for image_info in result["images"]:
        if os.path.exists(image_info["path"]):
            output_files.append({
                "path": image_info["path"],
                "name": image_info["filename"],
                "type": "image"
            })

    return output_files


def _format_extracted_text(result):
    """抽出されたテキストをフォーマット"""
    content = []

    content.append("=== PDF テキスト抽出結果 ===\n")

    if result["statistics"]:
        content.append("=== 抽出統計 ===")
        stats = result["statistics"]
        content.append(f"ファイルサイズ: {stats.get('file_size', 0):,} bytes")
        content.append(f"総ページ数: {stats.get('total_pages', 0)}")
        content.append(f"総文字数: {stats.get('total_characters', 0):,}")
        content.append(f"総単語数: {stats.get('total_words', 0):,}")
        content.append(f"総行数: {stats.get('total_lines', 0):,}")
        content.append(f"抽出方法: {stats.get('extraction_method', 'Unknown')}")
        content.append(f"キーワード一致数: {stats.get('keyword_matches', 0)}")
        content.append(f"抽出画像数: {stats.get('extracted_images', 0)}")
        content.append("")

    if result["keyword_matches"]:
        content.append("=== キーワード検索結果 ===")
        for i, match in enumerate(result["keyword_matches"], 1):
            content.append(f"[{i}] ページ {match['page']}, 行 {match['line_number']}")
            content.append(f"該当行: {match['line']}")
            content.append(f"前後の文脈:")
            content.append(match['context'])
            content.append("-" * 50)
        content.append("")

    content.append("=== 抽出されたテキスト ===")
    content.append(result["text"])

    return "\n".join(content)


def _create_zip_file(output_files, temp_dir):
    """複数ファイルをZIPファイルにまとめる"""
    zip_path = os.path.join(temp_dir, "extracted_content.zip")

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file_info in output_files:
            if os.path.exists(file_info["path"]):
                zipf.write(file_info["path"], file_info["name"])

    return zip_path


# ========================================
# 5. ページ回転機能（新機能）
# ========================================
@pdf_power_bp.route("/rotate", methods=["POST"])
@login_required
def rotate_pdf():
    """PDFページの回転機能"""
    temp_dir = None
    doc = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "PDFファイルのみ対応しています"}), 400

        rotation = request.form.get("rotation", "90")
        pages_str = request.form.get("pages", "all")

        try:
            rotation_angle = int(rotation)
            if rotation_angle not in [90, 180, 270]:
                return jsonify({"error": "回転角度は90, 180, 270のいずれかを指定してください"}), 400
        except ValueError:
            return jsonify({"error": "回転角度が正しくありません"}), 400

        temp_dir = tempfile.mkdtemp()
        input_path = _temp_input_pdf_path(temp_dir)
        output_path = os.path.join(temp_dir, "rotated_output.pdf")
        file.save(input_path)

        # ファイルサイズチェック
        if os.path.getsize(input_path) > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        doc = fitz.open(input_path)
        total_pages = len(doc)
        if total_pages == 0:
            return jsonify({"error": "PDFにページがありません"}), 400

        # ページ指定の解析
        try:
            pages_to_rotate = set(_parse_page_ranges(pages_str, total_pages))
        except ValueError:
            return jsonify({"error": "ページ指定が正しくありません（例: 1-3,5,7）"}), 400

        # ページを回転（既存回転を考慮して加算）
        for i in range(total_pages):
            if i not in pages_to_rotate:
                continue
            page = doc[i]
            current_rotation = int(getattr(page, "rotation", 0)) % 360
            page.set_rotation((current_rotation + rotation_angle) % 360)

        doc.save(output_path)
        doc.close()
        doc = None

        @after_this_request
        def cleanup(response):
            try:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except Exception as e:
                logger.error(f"一時ディレクトリの削除に失敗: {e}")
            return response

        return _send_path_bytes(output_path, "rotated_output.pdf", "application/pdf")

    except Exception as e:
        logger.error(f"回転処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"回転処理中にエラーが発生しました: {str(e)}"}), 500
    finally:
        if doc is not None:
            doc.close()


# ========================================
# 6. パスワード保護機能（新機能）
# ========================================
@pdf_power_bp.route("/password", methods=["POST"])
@login_required
def password_protect():
    """PDFにパスワード保護を追加/解除"""
    temp_dir = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "PDFファイルのみ対応しています"}), 400

        mode = request.form.get("mode", "add")
        password = request.form.get("password", "")

        if mode not in {"add", "remove"}:
            return jsonify({"error": "操作モードが正しくありません"}), 400

        if mode == "add" and not password:
            return jsonify({"error": "パスワードを入力してください"}), 400
        if mode == "remove" and not password:
            return jsonify({"error": "解除用のパスワードを入力してください"}), 400

        temp_dir = tempfile.mkdtemp()
        input_path = _temp_input_pdf_path(temp_dir)
        output_path = os.path.join(temp_dir, "password_protected.pdf")
        file.save(input_path)

        # ファイルサイズチェック
        if os.path.getsize(input_path) > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        try:
            if mode == "add":
                # パスワード追加（既に保護済みの場合は明示エラー）
                try:
                    with pikepdf.open(input_path) as pdf:
                        pdf.save(output_path, encryption=pikepdf.Encryption(
                            owner=password,
                            user=password,
                            R=6  # AES-256
                        ))
                except pikepdf.PasswordError:
                    return jsonify({"error": "このPDFには既にパスワードが設定されています。解除してから再設定してください"}), 400
            else:
                # パスワード解除（総当たりは行わず、入力パスワードのみ検証）
                try:
                    with pikepdf.open(input_path) as _:
                        return _render_pdf_power_form_error(
                            "このPDFにはパスワードが設定されていません。",
                            tab="password",
                            status=400,
                            password_mode="remove",
                            password_error="このPDFにはパスワードが設定されていません。",
                        )
                except pikepdf.PasswordError:
                    pass

                try:
                    with pikepdf.open(input_path, password=password) as pdf:
                        pdf.save(output_path)
                except pikepdf.PasswordError:
                    return _render_pdf_power_form_error(
                        "入力されたパスワードが正しくありません。もう一度確認してください。",
                        tab="password",
                        status=400,
                        password_mode="remove",
                        password_error="入力されたパスワードが正しくありません。もう一度確認してください。",
                    )

            @after_this_request
            def cleanup(response):
                try:
                    if temp_dir and os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.error(f"一時ディレクトリの削除に失敗: {e}")
                return response

            download_name = "password_protected.pdf" if mode == "add" else "password_removed.pdf"
            return _send_path_bytes(output_path, download_name, "application/pdf")

        except Exception as e:
            logger.error(f"パスワード処理に失敗しました: {e}")
            logger.error(traceback.format_exc())
            return jsonify({"error": f"パスワード処理に失敗しました: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"パスワード処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"パスワード処理中にエラーが発生しました: {str(e)}"}), 500
    finally:
        pass


# ========================================
# 7. メタデータ編集機能（新機能）
# ========================================
@pdf_power_bp.route("/metadata", methods=["POST"])
@login_required
def edit_metadata():
    """PDFメタデータの編集"""
    temp_dir = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "PDFファイルのみ対応しています"}), 400

        title = request.form.get("title", "")
        author = request.form.get("author", "")
        subject = request.form.get("subject", "")
        keywords = request.form.get("keywords", "")

        temp_dir = tempfile.mkdtemp()
        input_path = _temp_input_pdf_path(temp_dir)
        output_path = os.path.join(temp_dir, "metadata_updated.pdf")
        file.save(input_path)

        # ファイルサイズチェック
        if os.path.getsize(input_path) > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        try:
            doc = fitz.open(input_path)

            metadata = {}
            if title:
                metadata["title"] = title
            if author:
                metadata["author"] = author
            if subject:
                metadata["subject"] = subject
            if keywords:
                metadata["keywords"] = keywords

            doc.set_metadata(metadata)
            doc.save(output_path)
            doc.close()

            @after_this_request
            def cleanup(response):
                try:
                    if temp_dir and os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.error(f"一時ディレクトリの削除に失敗: {e}")
                return response

            return _send_path_bytes(output_path, "metadata_updated.pdf", "application/pdf")

        except Exception as e:
            logger.error(f"メタデータ編集に失敗しました: {e}")
            logger.error(traceback.format_exc())
            return jsonify({"error": f"メタデータ編集に失敗しました: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"メタデータ編集中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"メタデータ編集中にエラーが発生しました: {str(e)}"}), 500
    finally:
        pass


# ========================================
# 8. 透かし追加機能（新機能）
# ========================================
@pdf_power_bp.route("/watermark", methods=["POST"])
@login_required
def add_watermark():
    """PDFに透かしを追加"""
    temp_dir = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "PDFファイルのみ対応しています"}), 400

        watermark_text = request.form.get("watermark_text", "")
        if not watermark_text:
            return jsonify({"error": "透かしテキストを入力してください"}), 400

        temp_dir = tempfile.mkdtemp()
        input_path = _temp_input_pdf_path(temp_dir)
        output_path = os.path.join(temp_dir, "watermarked_output.pdf")
        file.save(input_path)

        # ファイルサイズチェック
        if os.path.getsize(input_path) > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        try:
            doc = fitz.open(input_path)

            for page_num in range(len(doc)):
                page = doc[page_num]
                overlay_pdf = fitz.open(
                    stream=_create_watermark_overlay(page.rect, watermark_text),
                    filetype="pdf",
                )
                try:
                    page.show_pdf_page(page.rect, overlay_pdf, 0, overlay=True)
                finally:
                    overlay_pdf.close()

            doc.save(output_path)
            doc.close()

            @after_this_request
            def cleanup(response):
                try:
                    if temp_dir and os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.error(f"一時ディレクトリの削除に失敗: {e}")
                return response

            return _send_path_bytes(output_path, "watermarked_output.pdf", "application/pdf")

        except Exception as e:
            logger.error(f"透かし追加に失敗しました: {e}")
            logger.error(traceback.format_exc())
            return jsonify({"error": f"透かし追加に失敗しました: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"透かし処理中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"透かし処理中にエラーが発生しました: {str(e)}"}), 500
    finally:
        pass


# ========================================
# 9. ページ操作機能（新機能）
# ========================================
@pdf_power_bp.route("/page_operations", methods=["POST"])
@login_required
def page_operations():
    """PDFページの削除・並べ替え・抽出"""
    temp_dir = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDFファイルがアップロードされていません"}), 400

        filename = _ensure_pdf_upload_filename(file, "uploaded.pdf")
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "PDFファイルのみ対応しています"}), 400

        operation = request.form.get("operation", "")
        pages_order = request.form.get("pages_order", "")

        if not operation or not pages_order:
            return jsonify({"error": "操作とページ指定が必要です"}), 400

        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, filename)
        output_path = os.path.join(temp_dir, "pages_output.pdf")
        file.save(input_path)

        # ファイルサイズチェック
        if os.path.getsize(input_path) > MAX_FILE_SIZE:
            return jsonify({"error": "ファイルサイズが大きすぎます（100MB以下にしてください）"}), 400

        try:
            reader = PdfReader(input_path)
            writer = PdfWriter()
            total_pages = len(reader.pages)

            # ページ順序の解析（例: "1,3,2,4-6"）
            page_indices = []
            for part in pages_order.split(','):
                part = part.strip()
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    page_indices.extend(range(start - 1, end))
                else:
                    page_indices.append(int(part) - 1)

            # ページの有効性チェック
            for idx in page_indices:
                if idx < 0 or idx >= total_pages:
                    return jsonify({"error": f"無効なページ番号: {idx + 1}（総ページ数: {total_pages}）"}), 400

            # 指定された順序でページを追加
            for idx in page_indices:
                writer.add_page(reader.pages[idx])

            with open(output_path, "wb") as output_pdf:
                writer.write(output_pdf)

            @after_this_request
            def cleanup(response):
                try:
                    if temp_dir and os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.error(f"一時ディレクトリの削除に失敗: {e}")
                return response

            return _send_path_bytes(output_path, "pages_output.pdf", "application/pdf")

        except ValueError as e:
            return jsonify({"error": f"ページ指定が正しくありません: {str(e)}"}), 400
        except Exception as e:
            logger.error(f"ページ操作に失敗しました: {e}")
            logger.error(traceback.format_exc())
            return jsonify({"error": f"ページ操作に失敗しました: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"ページ操作中にエラーが発生しました: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"ページ操作中にエラーが発生しました: {str(e)}"}), 500
    finally:
        pass


# ========================================
# ヘルスチェック用エンドポイント
# ========================================
@pdf_power_bp.route("/edit", methods=["POST"])
@login_required
def edit_pdf_overlay():
    """
    Overlay-based PDF editing.
    Supported operations:
    - add_text
    - add_shape (rect / ellipse / line)
    - add_image
    - add_freehand
    - copy_region_paste
    """
    temp_dir = None
    cleanup_registered = False
    doc = None
    reference_doc = None
    try:
        file = request.files.get("pdf")
        if not file or not file.filename:
            return jsonify({"error": "PDF file is required."}), 400
        if not _is_probably_pdf_upload(file):
            return jsonify({"error": "Only PDF files are supported."}), 400

        operations_raw = request.form.get("operations", "[]")
        try:
            operations = json.loads(operations_raw)
        except json.JSONDecodeError:
            return jsonify({"error": "Invalid operations JSON."}), 400

        if not isinstance(operations, list) or not operations:
            return jsonify({"error": "At least one edit operation is required."}), 400

        requires_reference_pdf = any(
            isinstance(op, dict)
            and op.get("type") == "copy_region_paste"
            and _normalize_edit_source_document(op.get("source_document")) == "reference"
            for op in operations
        )

        temp_dir = tempfile.mkdtemp()
        input_path = _temp_input_pdf_path(temp_dir)
        output_path = os.path.join(temp_dir, "powerpdf_edited_output.pdf")
        file.save(input_path)

        if os.path.getsize(input_path) > MAX_FILE_SIZE:
            return jsonify({"error": "File size exceeds 100MB limit."}), 400

        doc = fitz.open(input_path)
        total_pages = len(doc)
        if total_pages == 0:
            return jsonify({"error": "PDF has no pages."}), 400

        reference_total_pages = 0
        if requires_reference_pdf:
            reference_file = request.files.get("reference_pdf")
            if not reference_file or not reference_file.filename:
                return jsonify({"error": "Reference PDF is required for reference copy operations."}), 400
            if not _is_probably_pdf_upload(reference_file):
                return jsonify({"error": "Only PDF files are supported for reference_pdf."}), 400
            reference_input_path = os.path.join(temp_dir, "reference_input.pdf")
            reference_file.save(reference_input_path)
            if os.path.getsize(reference_input_path) > MAX_FILE_SIZE:
                return jsonify({"error": "Reference PDF size exceeds 100MB limit."}), 400
            reference_doc = fitz.open(reference_input_path)
            reference_total_pages = len(reference_doc)
            if reference_total_pages == 0:
                return jsonify({"error": "Reference PDF has no pages."}), 400

        affected_pages = _collect_affected_pages_for_edit(operations, total_pages, reference_total_pages)
        for page_idx in sorted(affected_pages["edit"]):
            page = doc[page_idx]
            if int(getattr(page, "rotation", 0)) % 360 != 0:
                page.remove_rotation()
        if reference_doc is not None:
            for page_idx in sorted(affected_pages["reference"]):
                page = reference_doc[page_idx]
                if int(getattr(page, "rotation", 0)) % 360 != 0:
                    page.remove_rotation()

        image_bytes_cache = {}
        for idx, op in enumerate(operations, start=1):
            if not isinstance(op, dict):
                return jsonify({"error": f"Operation {idx} format is invalid."}), 400

            op_type = op.get("type")
            if op_type == "add_text":
                page_idx = _parse_page_number(op.get("page"), total_pages)
                page = doc[page_idx]
                text = str(op.get("text", "")).strip()
                if not text:
                    raise ValueError(f"Operation {idx}: text is empty.")
                x = _round_to_tenth(op.get("x", 0))
                y = _round_to_tenth(op.get("y", 0))
                w = max(_round_to_tenth(op.get("width", 240)), 1)
                h = max(_round_to_tenth(op.get("height", 48)), 1)
                font_size = max(_round_to_tenth(op.get("font_size", 14)), 1)
                font_name = str(op.get("font_name", "helv"))
                color = _parse_color_hex(op.get("color", "#111827"), (0.07, 0.09, 0.13))
                _insert_textbox_with_fallback(
                    page=page,
                    rect=_to_pdf_rect(page, x, y, w, h),
                    text=text,
                    font_size=font_size,
                    color=color,
                    font_name=font_name,
                )

            elif op_type == "add_shape":
                page_idx = _parse_page_number(op.get("page"), total_pages)
                page = doc[page_idx]
                shape_type = str(op.get("shape", "rect")).lower()
                stroke_width = max(_round_to_tenth(op.get("stroke_width", 2)), 0.1)
                stroke = _parse_color_hex(op.get("stroke_color", "#1d4ed8"), (0.11, 0.31, 0.85))
                fill = _parse_color_hex(op.get("fill_color", "transparent"), None)
                dashes = _dash_pattern(str(op.get("stroke_style", "solid")).lower())
                shape = page.new_shape()
                if shape_type in {"rect", "ellipse"}:
                    rect = _to_pdf_rect(
                        page,
                        _round_to_tenth(op.get("x", 0)),
                        _round_to_tenth(op.get("y", 0)),
                        max(_round_to_tenth(op.get("width", 120)), 1),
                        max(_round_to_tenth(op.get("height", 80)), 1),
                    )
                    if shape_type == "rect":
                        shape.draw_rect(rect)
                    else:
                        shape.draw_oval(rect)
                elif shape_type == "line":
                    shape.draw_line(
                        _to_pdf_point(page, _round_to_tenth(op.get("x1", 0)), _round_to_tenth(op.get("y1", 0))),
                        _to_pdf_point(page, _round_to_tenth(op.get("x2", 120)), _round_to_tenth(op.get("y2", 120))),
                    )
                else:
                    raise ValueError(f"Operation {idx}: unsupported shape type '{shape_type}'.")
                shape.finish(color=stroke, fill=fill, width=stroke_width, dashes=dashes)
                shape.commit(overlay=True)

            elif op_type == "add_image":
                page_idx = _parse_page_number(op.get("page"), total_pages)
                page = doc[page_idx]
                image_ref = str(op.get("image_ref", "")).strip()
                image_bytes = b""
                if image_ref:
                    cache_key = f"ref:{image_ref}"
                    if cache_key not in image_bytes_cache:
                        image_file = request.files.get(image_ref)
                        if not image_file:
                            raise ValueError(f"Operation {idx}: image file not found for '{image_ref}'.")
                        image_bytes_cache[cache_key] = image_file.read()
                    image_bytes = image_bytes_cache[cache_key]
                    if not image_bytes:
                        raise ValueError(f"Operation {idx}: image data is empty for '{image_ref}'.")
                else:
                    image_files = request.files.getlist("image_files")
                    try:
                        image_index = int(op.get("image_index", 1))
                    except (TypeError, ValueError):
                        raise ValueError(f"Operation {idx}: image_index must be an integer.")
                    if image_index < 1 or image_index > len(image_files):
                        raise ValueError(f"Operation {idx}: image_index must be between 1 and {len(image_files)}.")
                    cache_key = f"index:{image_index}"
                    if cache_key not in image_bytes_cache:
                        image_bytes_cache[cache_key] = image_files[image_index - 1].read()
                    image_bytes = image_bytes_cache[cache_key]
                    if not image_bytes:
                        raise ValueError(f"Operation {idx}: image data is empty for image_index {image_index}.")
                page.insert_image(
                    _to_pdf_rect(
                        page,
                        _round_to_tenth(op.get("x", 0)),
                        _round_to_tenth(op.get("y", 0)),
                        max(_round_to_tenth(op.get("width", 120)), 1),
                        max(_round_to_tenth(op.get("height", 120)), 1),
                    ),
                    stream=image_bytes,
                    keep_proportion=False,
                    overlay=True,
                )

            elif op_type == "add_freehand":
                page_idx = _parse_page_number(op.get("page"), total_pages)
                page = doc[page_idx]
                points = op.get("points", [])
                if not isinstance(points, list) or len(points) < 2:
                    raise ValueError(f"Operation {idx}: freehand points must have at least 2 points.")
                pdf_points = []
                for point_idx, point in enumerate(points, start=1):
                    if not isinstance(point, dict):
                        raise ValueError(f"Operation {idx}: point {point_idx} is invalid.")
                    pdf_points.append(_to_pdf_point(page, _round_to_tenth(point.get("x", 0)), _round_to_tenth(point.get("y", 0))))
                shape = page.new_shape()
                for point_idx in range(1, len(pdf_points)):
                    shape.draw_line(pdf_points[point_idx - 1], pdf_points[point_idx])
                shape.finish(
                    color=_parse_color_hex(op.get("stroke_color", "#1d4ed8"), (0.11, 0.31, 0.85)),
                    fill=None,
                    width=max(_round_to_tenth(op.get("stroke_width", 2)), 0.1),
                    dashes=_dash_pattern(str(op.get("stroke_style", "solid")).lower()),
                )
                shape.commit(overlay=True)

            elif op_type == "copy_region_paste":
                source_document = _normalize_edit_source_document(op.get("source_document"))
                if source_document == "reference":
                    if reference_doc is None:
                        raise ValueError(f"Operation {idx}: reference PDF is required.")
                    source_page = reference_doc[_parse_page_number(op.get("source_page"), reference_total_pages)]
                else:
                    source_page = doc[_parse_page_number(op.get("source_page"), total_pages)]
                target_page = doc[_parse_page_number(op.get("target_page"), total_pages)]
                sx = _round_to_tenth(op.get("source_x", 0))
                sy = _round_to_tenth(op.get("source_y", 0))
                sw = max(_round_to_tenth(op.get("source_width", 120)), 1)
                sh = max(_round_to_tenth(op.get("source_height", 80)), 1)
                tx = _round_to_tenth(op.get("target_x", 0))
                ty = _round_to_tenth(op.get("target_y", 0))
                tw = max(_round_to_tenth(op.get("target_width", sw)), 1)
                th = max(_round_to_tenth(op.get("target_height", sh)), 1)
                clip = _to_pdf_rect(source_page, sx, sy, sw, sh)
                target_rect = _to_pdf_rect(target_page, tx, ty, tw, th)
                dpi_scale = COPY_REGION_BASE_DPI / 72.0
                pix = source_page.get_pixmap(
                    clip=clip,
                    matrix=fitz.Matrix(
                        min(COPY_REGION_MAX_SCALE, max(1.0, tw / sw) * dpi_scale),
                        min(COPY_REGION_MAX_SCALE, max(1.0, th / sh) * dpi_scale),
                    ),
                    alpha=False,
                )
                target_page.insert_image(target_rect, stream=pix.tobytes("png"), overlay=True)

            else:
                raise ValueError(f"Operation {idx}: unsupported type '{op_type}'.")

        doc.save(output_path)
        doc.close()
        doc = None
        if reference_doc is not None:
            reference_doc.close()
            reference_doc = None

        with open(output_path, "rb") as output_file:
            file_data = io.BytesIO(output_file.read())
        file_data.seek(0)

        @after_this_request
        def cleanup(response):
            try:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except Exception as e:
                logger.error(f"Temporary directory cleanup error: {e}")
            return response

        cleanup_registered = True
        return send_file(
            file_data,
            as_attachment=True,
            download_name="powerpdf_edited_output.pdf",
            mimetype="application/pdf",
        )

    except ValueError as e:
        logger.warning(f"PowerPDF edit validation error: {e}")
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        logger.error(f"PowerPDF edit error: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": f"PowerPDF edit failed: {str(e)}"}), 500
    finally:
        if doc is not None:
            doc.close()
        if reference_doc is not None:
            reference_doc.close()
        if temp_dir and not cleanup_registered and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)


@pdf_power_bp.route("/health", methods=["GET"])
def health_check():
    return jsonify({"status": "healthy", "service": "PowerPDF"}), 200
